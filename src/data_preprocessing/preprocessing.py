from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser

from PIL import Image
import io
import fitz  # PyMuPDF
from docx import Document as DocxDoc
import os
import uuid
import tempfile
import time
import math

from src.config.cloudinary import upload_image, get_image_url
from src.data_preprocessing.prompt import image_caption_prompt
from src.config.llm import llm_2_0 as llm
from src.config.vector_store import vector_store_lesson_content
from src.utils.logger import logger


def extract_and_chunk_documents(
    file_path: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    upload_images: bool = True,
    batch_size: int = 15,
):
    """
    1. Extract text and images from the document, keeping them in the order they appear.
    2. Upload images to Cloudinary and get captions using the image_caption_prompt in batches.
    3. Create separate chunks for text and images.

    Args:
        file_path: Path to the document file
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks
        upload_images: Whether to upload images to Cloudinary
        batch_size: Number of images to process in a single batch

    Returns:
        List of Document objects with separate text and image chunks
    """
    docs = []  # Store documents (both text and images)
    image_caption_chain = image_caption_prompt | llm | StrOutputParser()

    # Extract text and images from document
    if file_path.endswith(".docx"):
        docs = extract_docx_with_images(file_path)
    elif file_path.endswith(".pdf"):
        docs = extract_pdf_with_images(file_path)
    else:
        raise ValueError("Unsupported file type")

    # Separate text and image documents
    text_docs = [doc for doc in docs if doc.metadata.get("type") == "text"]
    image_docs = [doc for doc in docs if doc.metadata.get("type") == "image"]

    # Process images in batches: upload to Cloudinary and get captions
    processed_image_chunks = []
    if upload_images and image_docs:
        # Prepare image batches
        total_images = len(image_docs)
        num_batches = math.ceil(total_images / batch_size)

        logger.info(
            f"Processing {total_images} images in {num_batches} batches of size {batch_size}"
        )

        for batch_idx in range(num_batches):
            start_idx = batch_idx * batch_size
            end_idx = min((batch_idx + 1) * batch_size, total_images)
            current_batch = image_docs[start_idx:end_idx]

            logger.info(
                f"Processing batch {batch_idx+1}/{num_batches} with {len(current_batch)} images"
            )

            # Process each image in the batch (upload to Cloudinary)
            batch_image_data = []

            for doc in current_batch:
                if "image_data" in doc.metadata:
                    # Create a temporary file for the image
                    image_id = str(uuid.uuid4())
                    temp_dir = tempfile.mkdtemp()
                    img_path = os.path.join(temp_dir, f"{image_id}.png")

                    # Save image to temporary file
                    doc.metadata["image_data"].save(img_path)

                    # Upload to Cloudinary
                    upload_result = upload_image(
                        file_path=img_path,
                        folder="robokki_images",
                        public_id=image_id,
                    )

                    # Get public URL
                    public_url = upload_result["secure_url"]

                    # Store image data for batch processing
                    batch_image_data.append(
                        {
                            "public_url": public_url,
                            "temp_dir": temp_dir,
                            "img_path": img_path,
                        }
                    )

            # Process the batch with LLM (get captions)
            batch_inputs = []
            for img_data in batch_image_data:
                batch_inputs.append(
                    {
                        "messages": [
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "Hình này có chứa những gì vậy",
                                    },
                                    {
                                        "type": "image",
                                        "source_type": "url",
                                        "url": img_data["public_url"],
                                    },
                                ],
                            },
                        ],
                        "messages_history": [],
                    }
                )

            # Get captions for the batch
            try:
                batch_captions = image_caption_chain.batch(batch_inputs)

                # Create document chunks with captions
                for i, caption in enumerate(batch_captions):
                    processed_image_chunks.append(
                        Document(
                            page_content=caption,
                            metadata={
                                "type": "image",
                                "public_url": batch_image_data[i]["public_url"],
                            },
                        )
                    )

                    # Clean up temporary files
                    os.remove(batch_image_data[i]["img_path"])
                    os.rmdir(batch_image_data[i]["temp_dir"])

                # Sleep between batches to avoid rate limits if there are more batches
                if batch_idx < num_batches - 1:
                    sleep_time = 120  # 2 minutes between batches
                    logger.info(
                        f"Sleeping for {sleep_time} seconds before processing batch {batch_idx+2}/{num_batches}"
                    )
                    time.sleep(sleep_time)

            except Exception as e:
                logger.error(f"Error processing batch {batch_idx+1}: {str(e)}")
                # Clean up any remaining temporary files
                for img_data in batch_image_data:
                    try:
                        if os.path.exists(img_data["img_path"]):
                            os.remove(img_data["img_path"])
                        if os.path.exists(img_data["temp_dir"]):
                            os.rmdir(img_data["temp_dir"])
                    except Exception as cleanup_error:
                        logger.error(f"Error cleaning up: {str(cleanup_error)}")
                raise e

    # Process text documents - create a combined text document
    combined_text = ""
    for doc in text_docs:
        if combined_text:
            combined_text += "\n\n"
        combined_text += doc.page_content

    # Chunk the text
    text_chunks = []
    if combined_text:
        # Create a document with the combined text
        combined_doc = Document(page_content=combined_text, metadata={"type": "text"})

        # Split into chunks
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
        text_chunks = splitter.split_documents([combined_doc])

        # Ensure each text chunk has only the 'type' metadata
        for chunk in text_chunks:
            chunk.metadata = {"type": "text"}

    # Combine text chunks and image chunks in the original document order
    all_chunks = []
    text_idx, image_idx = 0, 0

    # Reconstruct the original order based on the input docs
    for doc in docs:
        if doc.metadata.get("type") == "text":
            if text_idx < len(text_chunks):
                all_chunks.append(text_chunks[text_idx])
                text_idx += 1
        elif doc.metadata.get("type") == "image":
            if image_idx < len(processed_image_chunks):
                all_chunks.append(processed_image_chunks[image_idx])
                image_idx += 1

    # Add any remaining chunks
    all_chunks.extend(text_chunks[text_idx:])
    all_chunks.extend(processed_image_chunks[image_idx:])

    return all_chunks


def extract_docx_with_images(path: str) -> list[Document]:
    """
    Extract text and images from DOCX file.

    Args:
        path: Path to the DOCX file

    Returns:
        List of Document objects containing text and images
    """
    doc = DocxDoc(path)
    docs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            docs.append(Document(page_content=text, metadata={"type": "text"}))

    for rel in doc.part._rels.values():
        if "image" in rel.target_ref:
            img_data = rel.target_part.blob
            image = Image.open(io.BytesIO(img_data))

            # Store image data in metadata for later processing
            docs.append(
                Document(
                    page_content="",  # Will be replaced with caption after processing
                    metadata={
                        "type": "image",
                        "image_data": image,
                    },
                )
            )

    return docs


def extract_pdf_with_images(pdf_path: str) -> list[Document]:
    """
    Extract text and images from PDF.

    Args:
        pdf_path: Path to the PDF file

    Returns:
        List of Document objects containing text and images
    """
    docs = []
    doc = fitz.open(pdf_path)

    # Extract text from PDF
    for page in doc:
        text = page.get_text("text")
        if text:
            docs.append(Document(page_content=text, metadata={"type": "text"}))

        # Extract images from PDF
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]

            # Convert image bytes to PIL Image
            image = Image.open(io.BytesIO(img_bytes))

            # Store image data in metadata for later processing
            docs.append(
                Document(
                    page_content="",  # Will be replaced with caption after processing
                    metadata={
                        "type": "image",
                        "image_data": image,
                    },
                )
            )

    return docs


def process_and_index_file(
    file_path: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    batch_size: int = 15,
) -> list[Document]:
    """
    Process a file and index it in the vector store.

    Args:
        file_path: Path to the file to process
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks
        batch_size: Number of images to process in a single batch

    Returns:
        List of processed Document objects
    """
    # Process the file
    documents = extract_and_chunk_documents(
        file_path=file_path,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        upload_images=True,
        batch_size=batch_size,
    )

    # Index in vector store
    # vector_store_lesson_content.add_documents(documents)

    return documents


def process_and_index_directory(
    directory_path: str,
    file_extensions: list[str] = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> list[Document]:
    """
    Process all files in a directory and index them in the vector store.

    Args:
        directory_path: Path to the directory
        file_extensions: List of file extensions to process (e.g., [".txt", ".md", ".pdf", ".docx"])
        chunk_size: Size of text chunks
        chunk_overlap: Overlap between chunks

    Returns:
        List of processed Document objects
    """
    all_docs = []

    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)

            # Skip files with unwanted extensions
            if file_extensions and not any(
                file.endswith(ext) for ext in file_extensions
            ):
                continue

            try:
                docs = process_and_index_file(
                    file_path=file_path,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                )
                all_docs.extend(docs)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")

    return all_docs


if __name__ == "__main__":
    # Example usage
    docs = process_and_index_file("./")
    print(f"Processed {len(docs)} chunks")

    # Or process a directory
    # docs = process_and_index_directory(
    #     "path/to/your/directory",
    #     file_extensions=[".txt", ".md", ".pdf", ".docx"]
    # )
    print(f"Processed {len(docs)} chunks from directory")
    pass
